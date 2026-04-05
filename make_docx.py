import os

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from PIL import Image
except ImportError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "Pillow"])
    from PIL import Image

with open("func_order.txt", "r", encoding="utf-8") as f:
    funcs = [line.strip() for line in f if line.strip() and line.strip() not in ("inputLine", "inputDate")]

doc = Document()

# Default styles for coursework (Times New Roman 14)
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

for i, func in enumerate(funcs, 1):
    img_path = os.path.join("out_charts_kod", f"{func}.png")
    caption_text = f"Рисунок 2.{i} – Блок-схема функции {func}"
    
    if os.path.exists(img_path):
        # Insert image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        
        # Calculate appropriate width
        with Image.open(img_path) as img:
            img_width, img_height = img.size
            # Assume 100 DPI for word document sizing
            desired_w_inches = img_width / 100.0
            
            # Cap width at 6.0 inches (typical max page width between margins)
            if desired_w_inches > 6.0:
                desired_w_inches = 6.0
                
        run.add_picture(img_path, width=Inches(desired_w_inches))
        
        # Insert caption underneath
        p_cap = doc.add_paragraph(caption_text)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # empty space between sections

# Save to the local working directory with a recognizable name
output_file = "Блок_схемы_Курсовая_v2.docx"
doc.save(output_file)
print(f"File {output_file} successfully generated.")
